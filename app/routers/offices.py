import json

import asyncpg
from fastapi import APIRouter, status, Depends
from fastapi.responses import FileResponse
from asyncpg import connect
from docx import Document
from sqlalchemy import select

from app.config import settings
from app.schemas.build import SMap, SMapPlace, SInventoryType
from app.schemas.office import SOfficeCreate, SFloorCreate, SOfficeInventory, SOfficeEmployee
from app.repositories.offices import OfficesRepository, FloorsRepository
from app.utils import get_admin_token
from app.db.base import async_session_maker
from app.db.models import Office


from app.utils import check_endpoint_permissions

router = APIRouter(
    prefix="/offices",
    tags=["Offices"],
    dependencies=[Depends(check_endpoint_permissions)]
)


@router.get('/offices', status_code=status.HTTP_200_OK)
async def get_offices():
    return await OfficesRepository.find_all()


@router.get('/floors/{office_id}', status_code=status.HTTP_200_OK)
async def get_office_floors(office_id: int):
    return await FloorsRepository.find_all(office_id=office_id)


@router.get('/inventory/{office_id}', status_code=status.HTTP_200_OK)
async def get_office_inventory(office_id: int) -> list[SOfficeInventory]:
    conn = await connect(settings.POSTGRES_CLEAR_URL)
    result = await conn.fetch(f"""
        SELECT user_inventory.id, inventory.name, users.fio FROM users
        JOIN user_inventory ON user_inventory.user_id = users.id
        JOIN inventory ON user_inventory.inventory_id = inventory.id
        WHERE office_id='{office_id}'
    """)
    return [SOfficeInventory(**elem) for elem in result]


@router.get('/employees/{office_id}')
async def get_office_employees(office_id) -> list[SOfficeEmployee]:
    conn = await connect(settings.POSTGRES_CLEAR_URL)
    await conn.set_type_codec(
        'json',
        encoder=json.dumps,
        decoder=json.loads,
        schema='pg_catalog'
    )
    result = await conn.fetch(f"""
        SELECT
            users.id,
            users.fio,
            users.position,
            users.email,
            coalesce(json_agg(json_build_object(
                'id', inventory.id, 'name', inventory.name
                )) filter (where inventory.id is not null), '[]'
            ) as inventory,
            coalesce(json_agg(json_build_object(
                'id', furniture.id
            )) filter (where furniture.id is not null), '[]'
            ) as furniture
        FROM users
        LEFT JOIN user_inventory ON user_inventory.user_id = users.id
        LEFT JOIN user_furniture ON user_furniture.user_id = users.id
        LEFT JOIN inventory ON user_inventory.inventory_id = inventory.id
        LEFT JOIN furniture ON user_furniture.furniture_id = furniture.id
        WHERE office_id='{office_id}'
        GROUP BY users.id
    """)
    return [SOfficeEmployee(**elem) for elem in result]


@router.get('/employees/{employee_id}/inventory')
async def get_employee_inventory(employee_id: str) -> list[SInventoryType]:
    conn = await connect(settings.POSTGRES_CLEAR_URL)
    result = await conn.fetch(f"""
        SELECT inventory.id, inventory.name FROM user_inventory
        JOIN inventory ON inventory.id = user_inventory.inventory_id
        WHERE user_id='{employee_id}'
    """)
    return [SInventoryType(**elem) for elem in result]


@router.get('/map/{office_id}')
async def get_map(office_id: int, floor_id: int | None = None) -> SMap:
    conn = await connect(settings.POSTGRES_CLEAR_URL)
    query = f"SELECT * FROM map WHERE office_id='{office_id}'"
    if floor_id:
        query += f"AND floor_id='{floor_id}'"
    result = await conn.fetch(query)
    return SMap(items=[SMapPlace(**item) for item in result])


@router.post("/office", status_code=status.HTTP_201_CREATED)
async def create_office(office: SOfficeCreate):
    return await OfficesRepository.create(**office.model_dump())


@router.post("/floor", status_code=status.HTTP_201_CREATED)
async def create_floor(floor: SFloorCreate):
    return await FloorsRepository.create(**floor.model_dump())


@router.put("/office/{office_id}", status_code=status.HTTP_204_NO_CONTENT)
async def update_office(
    office_id: int,
    office: SOfficeCreate
):
    await OfficesRepository.update(office_id, **office.model_dump())


@router.put("/floor/{floor_id}", status_code=status.HTTP_204_NO_CONTENT)
async def update_floor(
    floor_id: int,
    floor: SFloorCreate
):
    await OfficesRepository.update(floor_id, **floor.model_dump())


@router.delete("/office/{office_id}", status_code=status.HTTP_204_NO_CONTENT)
async def delete_office(office_id: int) -> None:
    await OfficesRepository.delete(id=office_id)


@router.delete("/floor/{floor_id}", status_code=status.HTTP_204_NO_CONTENT)
async def delete_floor(floor_id: int) -> None:
    await FloorsRepository.delete(floor_id)


@router.get('/stats', status_code=status.HTTP_200_OK)
async def get_statistics(access_token: str = Depends(get_admin_token)):
    doc = Document()
    conn: asyncpg.Connection = await connect(settings.POSTGRES_CLEAR_URL)
    await conn.set_type_codec(
        'json',
        encoder=json.dumps,
        decoder=json.loads,
        schema='pg_catalog'
    )

    doc.add_heading('Статистика по офисам компании', 0)
    offices = await conn.fetch("""
        SELECT 
            offices.name, 
            offices.address,
            count(DISTINCT floors.id) as floors_count,
            count(users.id) as users_count,
            count(user_inventory.id) as inventory_count
        FROM offices
        LEFT JOIN floors ON floors.office_id = offices.id
        LEFT JOIN users ON users.office_id = offices.id
        LEFT JOIN user_inventory ON users.id = user_inventory.user_id
        GROUP BY offices.id;
    """)
    doc.add_heading('Статистика по каждому офису', 1)
    for index, office in enumerate(offices):
        office = dict(office)
        doc.add_heading(f'Офис №{index+1}', 3)
        doc.add_paragraph(f'''
            Название: {office['name']}
            Адрес: {office['address']}
            
            Количество этажей: {office['floors_count']}
            Количество сотрудников: {office['users_count']}
            Количество инвентаря сотрудников: {office['inventory_count']}
        '''.replace('\n', '', 1))

    offices = await conn.fetchrow("""
        SELECT
            count(*) as offices_count,
            sum(users_count) as users_count,
            sum(floors_count) as floors_count,
            sum(users_count) / sum(floors_count) as users_per_floor,
            
            round(avg(users_count), 2) as avg_users,
            min(users_count) as min_users,
            max(users_count) as max_users,
            
            round(avg(floors_count), 2) as avg_floors,
            min(floors_count) as min_floors,
            max(floors_count) as max_floors
        FROM (
            SELECT 
                offices.name, 
                offices.address,
                count(DISTINCT floors.id) as floors_count,
                count(users.id) as users_count,
                count(user_inventory.id) as inventory_count
            FROM offices
            LEFT JOIN floors ON floors.office_id = offices.id
            LEFT JOIN users ON users.office_id = offices.id
            LEFT JOIN user_inventory ON users.id = user_inventory.user_id
            GROUP BY offices.id
        ) as idk
    """)
    doc.add_heading('Суммарная статистика по офисам', 1)
    doc.add_heading('Статистика по офисам', 3)
    doc.add_paragraph(f'''
        Всего офисов: {offices["offices_count"]}
        Всего этажей: {offices["floors_count"]}
        Всего сотрудников: {offices["users_count"]}
        В среднем сотрудников на этаже: {offices["users_per_floor"]}
    '''.replace('\n', '', 1))

    doc.add_heading('Статистика по сотрудникам в офисах', 3)
    doc.add_paragraph(f'''
        В среднем сотрудников в офисе: {offices["avg_users"]}
        Максимум сотрудников в офисе: {offices["max_users"]}
        Минимум сотрудников в офисе: {offices["min_users"]}
    '''.replace('\n', '', 1))

    doc.add_heading('Статистика по этажам в офисах', 3)
    doc.add_paragraph(f'''
        В среднем этажей в офисе: {offices["avg_floors"]}
        Максимум этажей в офисе: {offices["max_floors"]}
        Минимум этажей в офисе: {offices["min_floors"]}
    ''')

    doc.save('stats.docx')
    return FileResponse(path='stats.docx', filename='stats.docx', media_type='multipart/form-data')


    
